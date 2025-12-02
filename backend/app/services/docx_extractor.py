# import os
# import re
# import shutil
# import docx
# from docx.oxml import parse_xml
# from docx.opc.constants import RELATIONSHIP_TYPE as RT
# from paddleocr import PaddleOCR
# from PIL import Image
# from tqdm import tqdm

# ocr = PaddleOCR(lang='en')

# def ocr_image(image_path):
#     """Run OCR safely on image"""
#     try:
#         with Image.open(image_path) as img:
#             img.verify()
#         result = ocr.predict(image_path)
#         text_output = []
#         if result and len(result) > 0:
#             for line in result[0]:
#                 text_output.append(line[1][0])
#         return "\n".join(text_output)
#     except Exception as e:
#         print(f"Skipping unreadable image {image_path}: {e}")
#         return ""


# def extract_text_with_formatting_in_sequence(docx_path, image_dir="temp_docx_images_seq"):
#     doc = docx.Document(docx_path)
#     formatted_lines = []

#     if os.path.exists(image_dir):
#         shutil.rmtree(image_dir)
#     os.makedirs(image_dir, exist_ok=True)

#     rels = doc.part.rels
#     para_index = 0

#     for block in tqdm(doc.element.body.iterchildren(), desc="Extracting sequential content", ncols=90):
#         if block.tag.endswith("p"):
#             para = doc.paragraphs[para_index]
#             para_index += 1
#             text = para.text.strip()
#             if not text:
#                 continue

#             style_name = para.style.name.lower()

#             if "heading" in style_name:
#                 level = re.findall(r'\d+', style_name)
#                 level = int(level[0]) if level else 1
#                 prefix = "#" * level
#                 formatted_lines.append(f"\n{prefix} {text}\n")
#             elif para._element.xpath(".//w:numPr"):
#                 formatted_lines.append(f"- {text}")
#             else:
#                 formatted_lines.append(f"{text}")

#             # Inline images
#             for run in para.runs:
#                 if run.element.xpath(".//a:blip"):
#                     blip = run.element.xpath(".//a:blip")[0]
#                     embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
#                     if embed in rels:
#                         img_part = rels[embed].target_part
#                         img_data = img_part.blob
#                         img_name = f"image_seq_{len(formatted_lines)+1}.png"
#                         img_path = os.path.join(image_dir, img_name)
#                         with open(img_path, "wb") as f:
#                             f.write(img_data)
#                         ocr_text = ocr_image(img_path)
#                         if ocr_text.strip():
#                             formatted_lines.append(f"\n📷 **Extracted Text from Image:**\n> {ocr_text.strip()}\n")

#         elif block.tag.endswith("tbl"):
#             table_xml = parse_xml(block.xml)
#             temp_doc = docx.Document()
#             temp_doc._body.clear_content()
#             temp_doc._body._element.append(table_xml)
#             table = temp_doc.tables[0]

#             formatted_lines.append("\n---\n**Table:**\n")
#             for row in table.rows:
#                 cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
#                 formatted_lines.append("| " + " | ".join(cells) + " |")
#             formatted_lines.append("---\n")

#     formatted_text = "\n".join(formatted_lines)
#     formatted_text = re.sub(r'\n{3,}', '\n\n', formatted_text)

#     shutil.rmtree(image_dir, ignore_errors=True)
#     return formatted_text




import os
import re
import shutil
import docx
from docx.oxml import parse_xml
from paddleocr import PaddleOCR
from PIL import Image
from tqdm import tqdm

# Initialize Paddle OCR
ocr = PaddleOCR(lang='en')


def ocr_image(image_path):
    """Run OCR on image safely."""
    try:
        with Image.open(image_path) as img:
            img.verify()

        result = ocr.ocr(image_path, cls=False)
        text_output = []

        if result:
            for line in result:
                for text_block in line:
                    text_output.append(text_block[1][0])

        return "\n".join(text_output)

    except Exception as e:
        print(f"Skipping unreadable image {image_path}: {e}")
        return ""


def extract_text_with_formatting_in_sequence(docx_path, image_dir="temp_docx_images_seq"):
    """
    Extract structured text from DOCX including:
    - Headings (converted to markdown #)
    - Bullet lists
    - Inline images (OCR)
    - Tables
    - Paragraphs in exact sequential order
    """

    doc = docx.Document(docx_path)
    formatted_output = []

    # Reset image directory
    if os.path.exists(image_dir):
        shutil.rmtree(image_dir)
    os.makedirs(image_dir, exist_ok=True)

    rels = doc.part.rels  # relationship collection for embedded images
    para_index = 0

    for block in tqdm(doc.element.body.iterchildren(), desc=f"Extracting: {os.path.basename(docx_path)}", ncols=100):

        # ---------------------------------------------------
        # 1. PARAGRAPH
        # ---------------------------------------------------
        if block.tag.endswith("p"):
            paragraph = doc.paragraphs[para_index]
            para_index += 1

            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower()

            # Headings
            if "heading" in style_name:
                lvl = re.findall(r"\d+", style_name)
                lvl = int(lvl[0]) if lvl else 1
                formatted_output.append(f"\n{'#' * lvl} {text}\n")

            # Bulleted list
            elif paragraph._element.xpath(".//w:numPr"):
                formatted_output.append(f"- {text}")

            # Normal paragraph
            else:
                formatted_output.append(text)

            # Inline images inside runs
            for run in paragraph.runs:
                blips = run.element.xpath(".//a:blip")
                if not blips:
                    continue

                for blip in blips:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )

                    if not embed_id or embed_id not in rels:
                        continue

                    img_part = rels[embed_id].target_part
                    img_data = img_part.blob

                    img_name = f"inline_img_{len(formatted_output)}.png"
                    img_path = os.path.join(image_dir, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_data)

                    ocr_text = ocr_image(img_path)

                    if ocr_text.strip():
                        formatted_output.append(
                            f"\n📷 **Image Text:**\n> {ocr_text.strip()}\n"
                        )

        # ---------------------------------------------------
        # 2. TABLE
        # ---------------------------------------------------
        elif block.tag.endswith("tbl"):
            table_xml = parse_xml(block.xml)

            temp_doc = docx.Document()
            temp_doc._body.clear_content()
            temp_doc._body._element.append(table_xml)

            table = temp_doc.tables[0]

            formatted_output.append("\n---\n**Table:**\n")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                formatted_output.append("| " + " | ".join(cells) + " |")
            formatted_output.append("\n---\n")

    # Cleanup long empty spaces
    final_text = "\n".join(formatted_output)
    final_text = re.sub(r"\n{3,}", "\n\n", final_text)

    shutil.rmtree(image_dir, ignore_errors=True)

    return final_text
